import docx                                     #https://python-docx.readthedocs.io/en/latest/
import re 
import xlsxwriter                               #https://xlsxwriter.readthedocs.io/index.html
import cv2                                      #https://opencv-python-tutroals.readthedocs.io/en/latest/py_tutorials/py_tutorials.html
import fnmatch
import numpy as np
import pytesseract                              #https://pypi.org/project/pytesseract/
from wand.image import Image as wi              #http://docs.wand-py.org/en/0.5.7/
import os
import time
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\caio.santos\AppData\Local\Tesseract-OCR\tesseract.exe"


def extnum_from_DF(df):
    rex1 = re.compile('R\$\s?\d+?\,?\d?\d?\d?\.?\d?\d?\sm?b?illion')
    rex2 = re.compile('\d+?\.?\d?\d?\%')
    dic = []

    for t in df['Texto']:
        lis = []
        if rex1.search(t):

            for num in rex1.findall(t):
                lis.append(num)
        if rex2.search(t):
            for num in rex2.findall(t):
                lis.append(num)

        dic.append(lis)

    return df.join(pd.DataFrame(dic))


def doc_to_excel(filepath, excelname,padrão=None):
    doc = docx.Document(filepath)
    excel = xlsxwriter.Workbook(excelname + '.xlsx')
    sh = excel.add_worksheet()
    row = 0

    if padrão == None:
        rex1 = re.compile('R\$\s?\d+?\,?\d?\d?\d?\.?\d?\d?\sm?b?illion')
        rex2 = re.compile('\d+?\.?\d?\d?\%')
        for p in doc.paragraphs:
            col = 0

            if rex1.search(p.text):
                row += 1
                sh.write(row-1, col, p.text)


                for num in rex1.findall(p.text):
                    col += 1
                    sh.write(row-1, col, num)

                if rex2.search(p.text):
                    for num in rex2.findall(p.text):
                        col += 1
                        sh.write(row-1, col, num)

    else:
        if padrão == 'R$':
            rex = re.compile('R\$\s?\d+?\,?\d?\d?\d?\.?\d?\d?\sm?b?illion')

        if padrão == '%':
            rex = re.compile('\d+?\.?\d?\d?\%')

        for p in doc.paragraphs:
            col = 0
            if rex.search(p.text):
                row += 1
                sh.write(row-1, col, p.text)

                for num in rex.findall(p.text):
                    col += 1
                    sh.write(row-1, col, num)


    excel.close()

def table_to_excel(filepath, excelname, tableindex):
    doc = docx.Document(filepath)
    excel = xlsxwriter.Workbook(excelname + '.xlsx')
    r = 0

    if tableindex == 'all':
        for t in doc.tables:
            r = 0
            sh = excel.add_worksheet()
            for row in t.rows:
                r += 1
                c = 0
                for cel in row.cells:
                    for p in cel.paragraphs:
                        sh.write(r, c, p.text)
                        c += 1

        excel.close()
    else:
        t = doc.tables[tableindex]
        sh = excel.add_worksheet()
        for row in t.rows:
            r += 1
            c = 0
            for cel in row.cells:
                for p in cel.paragraphs:
                    sh.write(r, c, p.text)
                    c += 1

    excel.close()


def show_requestnumber(imagem):
    try:
        img = cv2.imread(imagem)
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)

        lower_range = np.array([0, 200, 200])
        upper_range = np.array([0, 255, 255])

        mask = cv2.inRange(hsv, lower_range, upper_range)

        _, bin = cv2.threshold(mask, 0, 255, cv2.THRESH_BINARY)
        cont, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)

        cv2.drawContours(img, cont, -1, (0, 255, 0), 3)
        for c in cont:
            (x, y, w, z) = cv2.boundingRect(c)
            roi = img[y:y + z, x:x + w]

        return img

    except TypeError:
        print("Não tem request nessa folha")


def get_requestnumbers_fromimg(imagem, isfile=True):
    try:
        if isfile == True:
            img = cv2.imread(imagem)
        else:
            img = imagem

        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        lower_range = np.array([0, 200, 200])
        upper_range = np.array([0, 255, 255])

        mask = cv2.inRange(hsv, lower_range, upper_range)

        _, bin = cv2.threshold(mask, 0, 255, cv2.THRESH_BINARY)
        cont, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)

        text = []

        for c in cont:
            (x, y, w, z) = cv2.boundingRect(c)
            roi = img[y:y + z, x:x + w]
            text.append(pytesseract.image_to_string(roi))

        return text
    except TypeError:
        print("Não tem request nessa folha")

def pdftojpgconvert(inputPath, filename, outputPath):
    os.chdir(inputPath)

    pdf = wi(filename=filename, resolution=300)
    pages = pdf.convert('jpeg')

    if outputPath == '':
        os.chdir(inputPath)
    else:
        os.chdir(outputPath)
    n = 0
    for img in pages.sequence:
        n += 1
        im = wi(img)
        im.save(filename='ima' + str(n) + '.jpg')
    # wand to save the images		Segundos total         em minutos	    em segundos por folha
    # 4 folhas: 4 segundos	                  4,50 	             0,08 	                     1,13
    # 20 folhas: 28 segundos	             28,00 	             0,47 	                     1,40
    # 79 folhas: 156.67 segundos            156,67 	             2,61 	                     1,98


def get_requesttoexcel(filepath, excelpath='', excelname='new workbook', pat=True):
    rex1 = re.compile('R\$?S?\s?\d+?\,?\d?\d?\d?\.?\d?\d?\sm?b?illion')
    rex2 = re.compile('R\$?S?\s?\d+?\,?\d?\d?\d?\.?\d?\d?\n')
    rex3 = re.compile('\d+?\.?\d?\d?\%')

    pdf = wi(filename=filepath, resolution=300)
    pdfimg = pdf.convert('jpeg')
    pages = []
    imgblobs = []

    for img in pdfimg.sequence:
        page=wi(image=img)
        imgblobs.append(page.make_blob('jpeg'))

    for im in imgblobs:
        npimg = np.asarray(bytearray(im), dtype=np.uint8)
        p = cv2.imdecode(npimg, cv2.IMREAD_UNCHANGED)
        pages.append(p)

    excel = xlsxwriter.Workbook(f'{excelpath}\\{excelname}.xlsx')
    sh = excel.add_worksheet('new sheet')

    r = 0
    for p in pages:
        c = 0
        r += 1
        page = pytesseract.image_to_string(p[3050:3500, 1000:1500])
        sh.write(r, c, f'n° Pagina OM: {page}')

        t = get_requestnumbers_fromimg(p, isfile=False)

        if pat == False:
            for p in t:
                c += 1
                sh.write(r, c, p)
        else:
            for p2 in t:
                if rex1.search(p2):
                    for num in rex1.findall(p2):
                        c += 1
                        sh.write(r, c, num)

                if rex2.search(p2):
                    for num in rex2.findall(p2):
                        c += 1
                        sh.write(r, c, num)

                if rex3.search(p2):
                    for num in rex3.findall(p2):
                        c += 1
                        sh.write(r, c, num)

    excel.close()
    #benchmaking com o filtro ativado:
    # For 4 pages: 16 seconds or 0.3 minutes - 4s/page
    # For 20 pages: 75 seconds or 1.2 minutes - 4s/page
    # For 78 pages: 365 seconds or 6.1 minutes - 5s/page

    # benchmaking sem o filtro ativado:
    # For 4 pages: 17 seconds or 0.3 minutes - 4s/page
    # For 20 pages: 73 seconds or 1.2 minutes - 4s/page
    # For 78 pages: 343 seconds or 5.7 minutes - 4s/page


    #benchmaking sem o filtro ativado

