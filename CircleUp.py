import docx
import re
import xlsxwriter
import cv2
import fnmatch
import numpy as np
import pytesseract
from wand.image import Image as wi
import os


pytesseract.pytesseract.tesseract_cmd = r"C:\Users\caio.santos\AppData\Local\Tesseract-OCR\tesseract.exe"

def doc_to_excel(filepath, padrão, excelname):
    doc = docx.Document(filepath)
    excel = xlsxwriter.Workbook(excelname + '.xlsx')
    sh = excel.add_worksheet()
    row = 0
    if padrão == 'R$':
        rex = re.compile('R\$\s?\d+?\,?\d?\d?\d?\.?\d?\d?\sm?b?illion')

    if padrão == '%':
        rex = re.compile('\d+?\.?\d?\d?\%')

    for p in doc.paragraphs:
        col = 0
        if rex.search(p.text):
            row += 1
            sh.write(row, col, p.text)
            col += 1
            for num in rex.findall(p.text):
                sh.write(row, col, num)
                col += 1

    excel.close()

def table_to_excel(filepath, excelname, tableindex):
    doc = docx.Document(filepath)
    excel = xlsxwriter.Workbook(excelname + '.xlsx')
    r = 0

    if tableindex == 'all':
        for t in doc.tables:
            r = 0
            sh = excel.add_worksheet()
            for row in t.rows:
                r += 1
                c = 0
                for cel in row.cells:
                    for p in cel.paragraphs:
                        sh.write(r, c, p.text)
                        c += 1

        excel.close()
    else:
        t = doc.tables[tableindex]
        sh = excel.add_worksheet()
        for row in t.rows:
            r += 1
            c = 0
            for cel in row.cells:
                for p in cel.paragraphs:
                    sh.write(r, c, p.text)
                    c += 1

    excel.close()

def show_requestnumber(imagem):
    try:
        img = cv2.imread(imagem)
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)

        lower_range = np.array([0, 200, 200])
        upper_range = np.array([0, 255, 255])

        mask = cv2.inRange(hsv, lower_range, upper_range)

        _, bin = cv2.threshold(mask, 0, 255, cv2.THRESH_BINARY)
        cont, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)

        cv2.drawContours(img, cont, -1, (0, 255, 0), 1)
        for c in cont:
            (x, y, w, z) = cv2.boundingRect(c)
            roi = img[y:y + z, x:x + w]
            cv2.imwrite('circ' + str(len(c))+ '.jpg', roi)


        cv2.imshow('Corner', img)
        cv2.waitKey()
    except TypeError:
        print("Não tem request nessa folha")

def get_requestnumbers_fromimg(imagem):
    try:
        img = cv2.imread(imagem)
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        lower_range = np.array([0, 200, 200])
        upper_range = np.array([0, 255, 255])

        mask = cv2.inRange(hsv, lower_range, upper_range)

        _, bin = cv2.threshold(mask, 0, 255, cv2.THRESH_BINARY)
        cont, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)

        text = []

        for c in cont:

            (x, y, w, z) = cv2.boundingRect(c)
            roi = img[y:y + z, x:x + w]
            text.append(pytesseract.image_to_string(roi))

        return text
    except TypeError:
        print("Não tem request nessa folha")

def pdftojpgconvert(inputPath, filename, outputPath):
    os.chdir(inputPath)

    pdf = wi(filename=filename, resolution=300)
    pages = pdf.convert('jpeg')

    if outputPath == '':
        os.chdir(inputPath)
    else:
        os.chdir(outputPath)
    n = 0
    for img in pages.sequence:
        n += 1
        im = wi(img)
        im.save(filename='ima' + str(n) + '.jpg')
    # wand to save the images		Segundos total         em minutos	    em segundos por folha
    # 4 folhas: 4 segundos	                  4,50 	             0,08 	                     1,13
    # 20 folhas: 28 segundos	             28,00 	             0,47 	                     1,40
    # 79 folhas: 156.67 segundos            156,67 	             2,61 	                     1,98


def get_requesttoexcel(inputpath, filename, outputpath):
    rex1 = re.compile('R\$?S?\s?\d+?\,?\d?\d?\d?\.?\d?\d?\sm?b?illion')
    rex2 = re.compile('R\$?S?\s?\d+?\,?\d?\d?\d?\.?\d?\d?\n')
    rex3 = re.compile('\d+?\.?\d?\d?\%')

    pdftojpgconvert(inputpath, filename, outputpath)

    pages = fnmatch.filter(os.listdir(), '*.jpg')

    excel = xlsxwriter.Workbook('new workbook.xlsx')
    sh = excel.add_worksheet('new sheet')

    r = 0

    for p in pages:
        c = 0
        r += 1
        page = pytesseract.image_to_string(cv2.imread(p)[3050:3500, 1000:1500])
        sh.write(r, c,'n° Pagina OM: ' + page)

        t = get_requestnumbers_fromimg(p)

        for p2 in t:
            if rex1.search(p2):
                for num in rex1.findall(p2):
                    c +=1
                    sh.write(r, c, num)

            if rex2.search(p2):
                for num in rex2.findall(p2):
                    c += 1
                    sh.write(r, c, num)

            if rex3.search(p2):
                for num in rex3.findall(p2):
                    c += 1
                    sh.write(r, c, num)

    excel.close()
    # 4 folhas: 15.2667   segundos	         15,27 	             0,25                   	 3,82
    # 20 folhas: 76.4620  segundos	         76,46 	             1,27 	                     3,82
    # 78 folhas: 333.7779 segundos	        333,78          	 5,56 	                     4,23


