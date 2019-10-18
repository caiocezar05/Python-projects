import docx
import os
import pandas as pd
import io
from PIL import Image
import pytesseract
from wand.image import Image as wi
import fnmatch
import re
import shutil
from sklearn.externals import joblib

pytesseract.pytesseract.tesseract_cmd = r"C:\Users\caio.santos\AppData\Local\Tesseract-OCR\tesseract.exe"


def get_files_fromdirectory(in_files, out_files):
    for root, dirs, files in os.walk(in_files):  # replace the . with your starting directory
        for file in files:
            path_file = os.path.join(root, file)
            shutil.copy2(path_file, out_files)  # change you destination dir


def clause_split(data):
    docnum = []
    clausula = []
    text = []

    rex = re.compile(r"cl[á,a]usula", flags=re.I)
    # rex = re.compile(r"(PARÁGRAFO\s\d+?º)|(\w+?:\d+?\.\d+?\.)")
    ndoc = 0

    if type(data) == str:
        cl = 0
        ndoc += 1
        for string in re.split(rex, data):
            cl += 1
            docnum.append(ndoc)
            clausula.append(cl - 1)
            text.append(string)

    else:
        for d in data:
            ndoc += 1
            cl = 0

            for string in re.split(rex, d):
                cl += 1
                docnum.append(ndoc)
                clausula.append(cl - 1)
                text.append(string)
    data = {'Doc': docnum, 'clausula': clausula, 'Text': text}
    df = pd.DataFrame(data=data)
    return df


def pdf2docx_convert(pdf_path, out):
    doc = docx.Document()
    text = get_pdf_text(pdf_path)
    doc.add_paragraph(text)
    doc.save(out + '.docx')


def get_pdf_text(pdf_path):
    pdf = wi(filename=pdf_path, resolution=300)
    pdfImg = pdf.convert('jpeg')
    imgBlobs = []
    text = []

    for img in pdfImg.sequence:
        page = wi(image=img)
        imgBlobs.append(page.make_blob('jpeg'))

    for imgBlob in imgBlobs:
        im = Image.open(io.BytesIO(imgBlob))
        tx = pytesseract.image_to_string(im, lang='por')
        text.append(tx)

    return text


def trainprepare(path):
    os.chdir(path)

    ndoc = 0
    docnum = []
    clausula = []
    text = []
    rex = re.compile(r"cl[á,a]usula", flags=re.I)
    for d in fnmatch.filter(os.listdir(), '*.docx'):
        doc = docx.Document(d)
        ndoc += 1
        cl = 0
        stringx = []
        for p in doc.paragraphs:
            stringx.append(p.text)

        for string in re.split(rex, ''.join(stringx)):
            cl += 1
            docnum.append(ndoc)
            clausula.append(cl - 1)
            text.append(string)

    data = {'Doc': docnum, 'clausula': clausula, 'Text': text}
    df = pd.DataFrame(data=data)
    return df


def trainprepare_DF_series(data):
    docnum = []
    clausula = []
    text = []

    rex = re.compile(r"cl[á,a]usula", flags=re.I)
    # rex = re.compile(r"(par[á,a]grafo\s\d+?º)|(\w+?:\d+?\.\d+?\.)")
    ndoc = 0
    for d in data:
        ndoc += 1
        cl = 0

        for string in re.split(rex, d):
            cl += 1
            docnum.append(ndoc)
            clausula.append(cl - 1)
            text.append(string)

    data = {'Doc': docnum, 'clausula': clausula, 'Text': text}
    df = pd.DataFrame(data=data)
    return df


def get_patter(df):
    text = df[0]
    cls = df[1]

    if cls == 'OBJ':
        string = []
        rex = re.compile(r"objeto...+\.")
        rex2 = re.compile(r"\n\n..+\.")
        for i in rex.findall(text):
            string.append(i)
        for i in rex2.findall(text):
            string.append(i)
        strin = '; '.join(string)

        return strin

    elif cls == 'VALOR':
        string = []
        rex = re.compile(r"R\$\s?\d.+?\s")
        for i in rex.findall(text):
            string.append(i)
        strin = '; '.join(string)

        return strin

    elif cls == 'MULTA':
        string = []
        rex = re.compile(r"\d+?\%.+\.")
        rex2 = re.compile(r"R\$\s?\d.+?\s")

        for i in rex.findall(text):
            string.append(i)

        for i in rex2.findall(text):
            string.append('; ')
            string.append(i)

        strin = ''.join(string)
        return strin


    elif cls == 'PRAZO':
        string = []
        rex = re.compile(r"\d\d\s?\(.+\)\s?\w+")
        rex2 = re.compile(r"\d\d\sde\s\w+\sde\s\d\d\d\d")
        rex3 = re.compile(r"\d\d\/\d\d\/\d\d\d\d")

        for i in rex.findall(text):
            string.append(i)

        for i in rex2.findall(text):
            string.append('; ')
            string.append(i)

        for i in rex3.findall(text):
            string.append('; ')
            string.append(i)

        strin = ''.join(string)
        return strin

    else:
        return "nothing"


def contract_reader_from_file(file):
    # model_adit_or_master = joblib.load('Master_aditiv_classifier')
    model_clause = joblib.load('Clause_classifier')

    if file.lower().endswith(('.pdf', '.PDF')):
        string = get_pdf_text(file)
        text = ''.join(string)


    elif file.endswith('.docx'):
        doc = docx.Document(file)
        string = []
        for p in doc.paragraphs:
            string.append(p.text)
        text = ''.join(string)

    df = clause_split(text)
    cl_class = model_clause.predict(df['Text'])
    # contr_type = model_adit_or_master.predict(text)

    # df['Contr Type'] = contr_type
    df['Cl class'] = cl_class

    df['Info'] = df[['Text', 'Cl class']].apply(get_patter, axis=1)
    return df


def contract_reader(data):
    model_clause = joblib.load('Clause_classifier')
    # model_adit_or_master = joblib.load('Master_aditiv_classifier')

    # contr_type = model_adit_or_master.predict(data)

    df = clause_split(data)
    cl_class = model_clause.predict(df['Text'])
    df['Cl class'] = cl_class
    # df['Contr Type'] = contr_type
    df['Info'] = df[['Text', 'Cl class']].apply(get_patter, axis=1)
    return df

os.chdir(r'C:\Users\caio.santos\Desktop\Allpark')
n=46

for f in fnmatch.filter(os.listdir(), '*.pdf')[n:]:
    n+=1
    pdf2docx_convert(f,f'{str(n)} - {f}')
    print(f)
    print(n)

